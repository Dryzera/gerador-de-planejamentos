from docx import Document
from docx.shared import Inches
import json
from sys import exit

document = Document()
font = document.styles['Normal'].font
font.name = 'Arial'

data_e_mes = input('Digite a data (ex: 23 de Maio): ')
dia_semana_str = input('Digite o dia da semana (ex: terça, quarta, quinta, sexta): ')

alunos_lista = ['Arthur Costa', 'Arthur Gabriel', 'Cristofer Fernandes', 'Kauã Henrique', 'Mariana Carolina', 'Nicolas Kenji', 'Jean Davi']
dia_semana = {}

PATH_DEFAULT = '.\\database_json\\'

def load_database_aulas(day_obtido, path_padrao):
    path_padrao += f'{day_obtido}.json'
    with open(path_padrao, 'r', encoding='utf8') as file:
        day_week = json.load(file)
    return day_week

def catch_day_week(day):
    if day == 'terça':
        day_week = load_database_aulas('terca_feira', PATH_DEFAULT)
        return day_week
    elif day == 'quarta':
        day_week = load_database_aulas('quarta_feira', PATH_DEFAULT)
        return day_week
    elif day == 'quinta':
        day_week =load_database_aulas ('quinta_feira', PATH_DEFAULT)
        return day_week
    elif day == 'sexta':
        day_week = load_database_aulas('sexta_feira', PATH_DEFAULT)
        return day_week
    else:
        print(f'Você me enviou "{day}", aceito apenas: [terça, quarta, quinta ou sexta]')
        exit()

def criar_corpo_documento(command_data, day, db, alunos):
    numero_da_aula = 1
    paragrafo = document.add_paragraph(f'Londrina, {command_data} - {day.capitalize()}-feira\n')
    paragrafo.add_run('\n')
    lista_de_aulas = []

    while numero_da_aula <= 5:
        aula_iterada = db[f'Aula {numero_da_aula}']
        if aula_iterada in ['Matemática', 'Português', 'Ciências'] and aula_iterada not in lista_de_aulas:
            paragrafo.add_run(f'{numero_da_aula}° Aula: {aula_iterada}\n').bold = True
            paragrafo.add_run('Conteúdo: \n').bold = True
            paragrafo.add_run('Objetivo: \n').bold = True
            paragrafo.add_run('Metodologia: \n').bold = True
            paragrafo.add_run('\n')
            numero_da_aula += 1
            lista_de_aulas.append(aula_iterada)

        elif aula_iterada in lista_de_aulas:
            paragrafo.add_run(f'{numero_da_aula}° Aula: {aula_iterada}\n').bold = True
            paragrafo.add_run('continuação... \n').bold = True
            paragrafo.add_run('\n')
            numero_da_aula += 1
        else:
            paragrafo.add_run(f'{numero_da_aula}° Aula: {aula_iterada}\n').bold = True
            paragrafo.add_run('\n')
            numero_da_aula += 1

    if numero_da_aula == 6:
        for aluno in alunos:
            paragrafo.add_run(aluno).bold = True
            paragrafo.add_run('\n')
        salvar_aula()


def salvar_aula():
    document.add_page_break()
    document.save(f'.\\planejamentos\\Londrina {data_e_mes} {dia_semana_str}-feira.docx')
    print('Aula criada com sucesso!')
    exit()

if data_e_mes and dia_semana_str is not None:
    dia_semana = catch_day_week(dia_semana_str)
    criar_corpo_documento(data_e_mes, dia_semana_str, dia_semana, alunos_lista)
else:
    print('Você não me enviou algum valor.')